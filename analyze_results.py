import pandas as pd
import matplotlib.pyplot as plt

from pathlib import Path
from docx import Document
from docx.shared import Inches

CSV = Path("eval_runs.csv")
if not CSV.exists():
    raise SystemExit("eval_runs.csv not found. Run runner_macos.py first.")

df = pd.read_csv(CSV)

def agg_metrics(frame):
    acc = frame["success"].mean() * 100.0 if len(frame) else float('nan')
    lat = frame.loc[frame["success"] == 1, "latency_ms"].mean() if (frame["success"] == 1).any() else float('nan')
    eff = frame["effort_saved_pct"].mean() if "effort_saved_pct" in frame and len(frame) else float('nan')
    return pd.Series({"Accuracy_%": round(acc,2), "Latency_ms": round(lat,1) if pd.notna(lat) else float('nan'),
                      "EffortSaved_%": round(eff,1) if pd.notna(eff) else float('nan')})

overall = agg_metrics(df)
by_test = df.groupby("test_case").apply(agg_metrics).reset_index()

para = df[df["paraphrase_of"].notna() & (df["paraphrase_of"] != "")]
rob_overall = round(para["success"].mean()*100.0, 2) if len(para) else float('nan')
rob_by_test = para.groupby("test_case")["success"].mean().mul(100.0).round(2).reset_index().rename(columns={"success":"Robustness_%"})

def barplot(series, title, ylabel, out):
    plt.figure(figsize=(6,4))
    plt.bar(series.index, series.values)
    plt.title(title); plt.xlabel("Test case"); plt.ylabel(ylabel)
    plt.tight_layout(); plt.savefig(out); plt.close()

barplot(by_test.set_index("test_case")["Accuracy_%"], "Accuracy per test case", "Accuracy (%)", "plot_accuracy.png")
barplot(by_test.set_index("test_case")["Latency_ms"], "Latency per test case", "Latency (ms)", "plot_latency.png")
barplot(by_test.set_index("test_case")["EffortSaved_%"], "Effort saved per test case", "Effort saved (%)", "plot_effort.png")
barplot(rob_by_test.set_index("test_case")["Robustness_%"], "Robustness per test case", "Robustness (%)", "plot_robustness.png")

doc = Document()
doc.add_heading("Experimental Results", level=1)

doc.add_heading("Overall metrics", level=2)
t = doc.add_table(rows=2, cols=3)
hdr = t.rows[0].cells; hdr[0].text="Accuracy (%)"; hdr[1].text="Latency (ms)"; hdr[2].text="Effort Saved (%)"
row = t.rows[1].cells; row[0].text=str(overall["Accuracy_%"]); row[1].text=str(overall["Latency_ms"]); row[2].text=str(overall["EffortSaved_%"])

doc.add_heading("Per-test metrics", level=2)
t2 = doc.add_table(rows=1, cols=4)
hdr = t2.rows[0].cells; hdr[0].text="Test case"; hdr[1].text="Accuracy (%)"; hdr[2].text="Latency (ms)"; hdr[3].text="Effort Saved (%)"
for _, r in by_test.iterrows():
    cells = t2.add_row().cells
    cells[0].text = r["test_case"]
    cells[1].text = str(r["Accuracy_%"])
    cells[2].text = str(r["Latency_ms"])
    cells[3].text = str(r["EffortSaved_%"])

doc.add_heading("Robustness", level=2)
doc.add_paragraph(f"Overall robustness on paraphrases: {rob_overall} %")

doc.add_heading("Figures", level=2)
for p in ["plot_accuracy.png","plot_latency.png","plot_effort.png","plot_robustness.png"]:
    if Path(p).exists():
        doc.add_picture(p, width=Inches(5.5))

doc.save("Results_Report.docx")
print("Saved Results_Report.docx")
